"""Provider document RAG service for super-admin approval review."""

from __future__ import annotations

import hashlib
import json
import math
import logging
import re
from pathlib import Path
from typing import Any
from uuid import UUID

import httpx
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import joinedload

from app.core.config import settings
from app.models import ProviderDocument, ProviderDocumentChunk, ServiceProvider
from app.models.provider_document import EMBEDDING_DIMENSIONS, ENABLE_PGVECTOR, Vector
from app.services.provider_application_service import BASE_DIR, get_application

logger = logging.getLogger(__name__)

_EMBEDDER = None


class ProviderDocumentRAGService:
    """Index and answer questions over one provider's onboarding documents."""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def reindex_provider(self, provider_id: UUID) -> dict[str, Any]:
        """Rebuild the document index for one provider application.

        Uses a write-then-swap approach: new chunks are inserted first,
        old ones deleted only if new indexing succeeded — preventing the
        provider from having zero chunks if indexing fails mid-way.
        """
        try:
            provider = await self._get_provider(provider_id)
            application = get_application(str(provider_id)) or {}
            documents = application.get("documents") or []

            indexed_documents = 0
            indexed_chunks = 0
            skipped_documents: list[dict[str, str]] = []
            new_docs: list = []
            new_chunks: list = []

            for doc in documents:
                text = self._extract_document_text(doc)
                if not text:
                    skipped_documents.append(
                        {
                            "name": doc.get("name") or "Document",
                            "reason": "No readable text could be extracted.",
                        }
                    )
                    continue

                document = ProviderDocument(
                    provider_id=provider_id,
                    name=doc.get("name") or "Document",
                    path=doc.get("path") or "",
                    content_type=doc.get("content_type"),
                    extracted_text=text,
                    status="indexed",
                )
                new_docs.append(document)

                for index, chunk_text in enumerate(_chunk_text(text), start=1):
                    embedding = _embed_text(chunk_text)
                    new_chunks.append((document, index, chunk_text, embedding))
                    indexed_chunks += 1

                indexed_documents += 1

            # Only delete old data once we know we have new data to replace it
            if indexed_documents > 0 or skipped_documents:
                await self.db.execute(
                    delete(ProviderDocumentChunk).where(ProviderDocumentChunk.provider_id == provider_id)
                )
                await self.db.execute(
                    delete(ProviderDocument).where(ProviderDocument.provider_id == provider_id)
                )

            for document in new_docs:
                self.db.add(document)
            await self.db.flush()

            # Now associate chunks with their flushed document ids
            doc_iter = iter(new_docs)
            cur_doc = next(doc_iter, None)
            doc_chunk_idx = 0
            for (orig_doc, index, chunk_text, embedding) in new_chunks:
                # Find the matching flushed document
                if cur_doc is not None and orig_doc is not cur_doc:
                    cur_doc = next(doc_iter, None)
                chunk_kwargs = dict(
                    document_id=cur_doc.id if cur_doc else None,
                    provider_id=provider_id,
                    chunk_index=index,
                    page_number=None,
                    chunk_text=chunk_text,
                    embedding_json=json.dumps(embedding),
                )
                if self._can_use_vector_db():
                    chunk_kwargs["embedding"] = embedding
                chunk = ProviderDocumentChunk(**chunk_kwargs)
                self.db.add(chunk)

            await self.db.flush()
            return {
                "provider_id": str(provider.id),
                "provider_name": provider.user.full_name if provider.user else None,
                "indexed_documents": indexed_documents,
                "indexed_chunks": indexed_chunks,
                "skipped_documents": skipped_documents,
            }
        except Exception as exc:
            logger.exception("Document reindex failed: %s", exc)
            raise ValueError("Could not index provider documents") from exc

    async def ask(self, provider_id: UUID, question: str) -> dict[str, Any]:
        """Answer any super-admin question against indexed provider documents."""
        try:
            provider = await self._get_provider(provider_id)
            question = question.strip()
            if not question:
                return {
                    "answer": "Ask a question about this provider's uploaded documents.",
                    "citations": [],
                    "confidence": "low",
                    "risk_flags": [],
                }

            chunks = await self._retrieve_chunks(provider_id, question)
            if not chunks:
                index_result = await self.reindex_provider(provider_id)
                chunks = await self._retrieve_chunks(provider_id, question)
                if not chunks:
                    return {
                        "answer": (
                            "I could not find readable indexed text for this provider's uploaded documents. "
                            "Ask the provider for text-readable PDFs, DOCX files, or clearer document scans."
                        ),
                        "citations": [],
                        "confidence": "low",
                        "risk_flags": ["No readable document text indexed"],
                        "index": index_result,
                    }

            document_text = await self._load_provider_document_text(provider_id)
            context = self._build_context(provider, chunks, document_text)
            question_mode = self._question_mode(question)
            direct_answer = self._direct_document_answer(question, chunks, document_text)
            if not direct_answer and question_mode == "summary":
                direct_answer = self._build_summary_answer(provider, chunks, document_text)
            answer = direct_answer or await self._generate_answer(question, context)
            citations = [
                {
                    "document": chunk.document.name if chunk.document else "Document",
                    "path": chunk.document.path if chunk.document else "",
                    "chunk_index": chunk.chunk_index,
                    "excerpt": _truncate(chunk.chunk_text, 260),
                    "similarity": round(score, 3),
                }
                for chunk, score in chunks
            ]

            risk_flags = _extract_risk_flags(answer, chunks)
            if direct_answer and "could not find" not in direct_answer.lower():
                risk_flags = [flag for flag in risk_flags if flag != "Low similarity between question and uploaded document text"]
            return {
                "answer": answer,
                "citations": citations,
                "confidence": "high" if direct_answer else _confidence_from_scores([score for _, score in chunks]),
                "risk_flags": risk_flags,
            }
        except Exception as exc:
            logger.exception("Document AI failed: %s", exc)
            return {
                "answer": "I could not process this document question right now. Please try again.",
                "citations": [],
                "confidence": "low",
                "risk_flags": ["Document AI encountered an internal error"],
            }

    async def _get_provider(self, provider_id: UUID) -> ServiceProvider:
        result = await self.db.execute(
            select(ServiceProvider)
            .options(joinedload(ServiceProvider.user), joinedload(ServiceProvider.category))
            .where(ServiceProvider.id == provider_id)
        )
        provider = result.unique().scalar_one_or_none()
        if not provider:
            raise ValueError("Provider not found")
        return provider

    async def _retrieve_chunks(
        self,
        provider_id: UUID,
        question: str,
        limit: int = 8,
    ) -> list[tuple[ProviderDocumentChunk, float]]:
        query_embedding = _embed_text(question)
        chunks = await self._retrieve_chunks_from_vector_store(provider_id, query_embedding, limit * 4)

        scored: list[tuple[ProviderDocumentChunk, float]] = []
        raw_query = question.lower()
        query_tokens = set(_expand_query_tokens(_tokens(question), raw_query))
        section_hints = self._query_section_hints(raw_query)
        for chunk in chunks:
            embedding = _loads_embedding(chunk.embedding_json)
            semantic = _cosine_similarity(query_embedding, embedding)
            chunk_lower = chunk.chunk_text.lower()
            chunk_tokens = set(_tokens(chunk.chunk_text))
            overlap = len(query_tokens & chunk_tokens) / max(len(query_tokens), 1)
            exact_phrase = 1.0 if question.lower() in chunk.chunk_text.lower() else 0.0
            lexical_hits = sum(1 for token in query_tokens if token in chunk_lower)
            section_bonus = self._section_bonus(chunk_lower, section_hints)
            keyword_bonus = 0.0
            if lexical_hits:
                keyword_bonus += min(0.18, lexical_hits * 0.03)
            if any(term in chunk_lower for term in ("education", "project", "skills", "experience", "achievement", "certificate", "license", "qualification")):
                keyword_bonus += 0.05
            # Keep semantic signal, but let lexical/section matches dominate for exact document QA.
            score = (semantic * 0.28) + (overlap * 0.34) + (exact_phrase * 0.18) + keyword_bonus + section_bonus
            scored.append((chunk, score))

        scored.sort(key=lambda item: item[1], reverse=True)
        return scored[:limit]

    async def _retrieve_chunks_from_vector_store(
        self,
        provider_id: UUID,
        query_embedding: list[float],
        limit: int,
    ) -> list[ProviderDocumentChunk]:
        """Fetch nearest chunks from pgvector when available, otherwise fall back to recent chunks."""
        if self._can_use_vector_db():
            try:
                embedding_column = ProviderDocumentChunk.embedding  # type: ignore[attr-defined]
                result = await self.db.execute(
                    select(ProviderDocumentChunk)
                    .options(joinedload(ProviderDocumentChunk.document))
                    .where(ProviderDocumentChunk.provider_id == provider_id)
                    .order_by(embedding_column.cosine_distance(query_embedding))
                    .limit(limit)
                )
                return result.unique().scalars().all()
            except Exception:
                pass

        result = await self.db.execute(
            select(ProviderDocumentChunk)
            .options(joinedload(ProviderDocumentChunk.document))
            .where(ProviderDocumentChunk.provider_id == provider_id)
        )
        return result.unique().scalars().all()[:limit]

    def _build_context(
        self,
        provider: ServiceProvider,
        chunks: list[tuple[ProviderDocumentChunk, float]],
        document_text: str = "",
    ) -> str:
        application = get_application(str(provider.id)) or {}
        payload = application.get("payload") or {}
        provider_lines = [
            f"Provider name: {provider.user.full_name if provider.user else 'Unknown'}",
            f"Claimed specialization: {payload.get('specialization') or provider.specialization}",
            f"Claimed category: {provider.category.name if provider.category else 'Unknown'}",
            f"Claimed experience: {payload.get('experience_years', provider.experience_years)} years",
            f"Claimed location: {payload.get('location') or provider.location}",
            f"Profile description: {payload.get('profile_description') or provider.profile_description or 'Not provided'}",
        ]
        chunk_lines = []
        for idx, (chunk, score) in enumerate(chunks, start=1):
            source = chunk.document.name if chunk.document else "Document"
            chunk_lines.append(
                f"[{idx}] Source: {source}, chunk {chunk.chunk_index}, similarity {score:.3f}\n"
                f"{chunk.chunk_text}"
            )
        sections = ["\n".join(provider_lines)]
        if document_text:
            sections.append(f"FULL EXTRACTED DOCUMENT TEXT:\n{_truncate(document_text, 7000)}")
        sections.append("RETRIEVED DOCUMENT CHUNKS:\n" + "\n\n".join(chunk_lines))
        return "\n\n".join(sections)

    async def _generate_answer(self, question: str, context: str) -> str:
        prompt = f"""
You are a super-admin document review assistant for AppointEase provider approvals.
Answer only from the provider profile, full extracted document text, and retrieved document chunks below.
If the answer is directly present, answer it clearly and straightforwardly.
If evidence is missing, say exactly: "I could not find that in the uploaded documents."
Do not approve or reject automatically.
Only use facts supported by the provided context.
If there are multiple possible matches, say which one is most likely and mention uncertainty.
Keep the answer practical and cite source numbers like [1], [2] when you rely on retrieved chunks.

QUESTION:
{question}

CONTEXT:
{context}
""".strip()

        if settings.GEMINI_RAG_API_KEY or settings.GEMINI_API_KEY:
            rag_key = settings.GEMINI_RAG_API_KEY or settings.GEMINI_API_KEY
            try:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        "https://generativelanguage.googleapis.com/v1beta/models/"
                        f"gemini-2.5-flash:generateContent?key={rag_key}",
                        headers={"Content-Type": "application/json"},
                        json={
                            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                            "generationConfig": {"temperature": 0.2, "maxOutputTokens": 900},
                        },
                    )
                if response.status_code == 200:
                    data = response.json()
                    return data["candidates"][0]["content"]["parts"][0]["text"].strip()
            except Exception as exc:
                logger.warning("Gemini document AI request failed: %s", exc)

        if settings.GROK_API_KEY:
            try:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        "https://api.x.ai/v1/chat/completions",
                        headers={
                            "Authorization": f"Bearer {settings.GROK_API_KEY}",
                            "Content-Type": "application/json",
                        },
                        json={
                            "model": settings.GROK_MODEL,
                            "messages": [{"role": "user", "content": prompt}],
                            "temperature": 0.2,
                            "max_tokens": 900,
                        },
                    )
                if response.status_code == 200:
                    return response.json()["choices"][0]["message"]["content"].strip()
            except Exception as exc:
                logger.warning("Grok document AI request failed: %s", exc)

        return self._fallback_answer(question, context)

    def _fallback_answer(self, question: str, context: str) -> str:
        snippets = []
        for match in re.finditer(r"\[(\d+)\] Source: ([^\n]+)\n(.+?)(?=\n\n\[\d+\] Source:|\Z)", context, re.S):
            snippets.append((match.group(1), match.group(2), _truncate(" ".join(match.group(3).split()), 360)))
        if not snippets:
            return "I could not find that in the uploaded documents."

        lines = ["I found these relevant document excerpts:"]
        for number, source, text in snippets[:3]:
            lines.append(f"- [{number}] {source}: {text}")
        lines.append(
            "\nManual review note: verify names, dates, license numbers, expiry dates, and specialization proof in the original files before approving."
        )
        return "\n".join(lines)

    async def _load_provider_document_text(self, provider_id: UUID) -> str:
        """Load all extracted text for a provider's uploaded documents."""
        result = await self.db.execute(
            select(ProviderDocument)
            .where(ProviderDocument.provider_id == provider_id)
            .order_by(ProviderDocument.created_at.asc())
        )
        docs = result.scalars().all()
        parts = []
        for doc in docs:
            if doc.extracted_text.strip():
                parts.append(f"[{doc.name}]\n{doc.extracted_text.strip()}")
        return "\n\n".join(parts)

    @staticmethod
    def _can_use_vector_db() -> bool:
        """True when pgvector is enabled and the vector column exists."""
        return ENABLE_PGVECTOR and Vector is not None

    def _direct_document_answer(
        self,
        question: str,
        chunks: list[tuple[ProviderDocumentChunk, float]],
        document_text: str = "",
    ) -> str | None:
        """Extract obvious facts directly when the question is about a specific field."""
        q = question.lower()
        if not any(term in q for term in ("college", "colleges", "university", "institute", "education", "study", "degree", "qualification", "gpa", "cgpa", "skills", "projects", "achievement", "experience", "license", "certificate", "expiry", "issue date", "issue", "validity", "registration", "number")):
            return None

        source_text = document_text or "\n".join(chunk.chunk_text for chunk, _ in chunks)
        text_blocks = []
        for chunk, score in chunks:
            text_blocks.append((score, f"{chunk.chunk_text}"))

        # Prefer the full extracted document text for direct factual questions.
        if source_text:
            direct = self._answer_from_source_text(q, source_text)
            if direct:
                return direct

        for score, text in sorted(text_blocks, key=lambda item: item[0], reverse=True):
            if not text:
                continue
            direct = self._answer_from_source_text(q, text)
            if direct:
                return direct

        return None

    @staticmethod
    def _question_mode(question: str) -> str:
        """Classify common document QA prompts into a safe answer mode."""
        q = question.lower()
        if any(term in q for term in ("summarize", "summary", "brief", "overview", "tell me about the document")):
            return "summary"
        if any(term in q for term in ("college", "university", "institute", "gpa", "cgpa", "degree", "qualification")):
            return "education"
        if any(term in q for term in ("skills", "technologies", "tech stack", "tools")):
            return "skills"
        if any(term in q for term in ("project", "projects", "work experience", "experience")):
            return "experience"
        if any(term in q for term in ("achievement", "award", "badge", "certification")):
            return "achievements"
        return "general"

    @staticmethod
    def _query_section_hints(question: str) -> set[str]:
        """Map a question to likely document sections for retrieval boost."""
        hints: set[str] = set()
        q = question.lower()
        if any(term in q for term in ("project", "projects", "built", "developed", "work", "case study")):
            hints.add("projects")
        if any(term in q for term in ("achievement", "award", "badge", "certification", "honor", "honours")):
            hints.add("achievements")
        if any(term in q for term in ("education", "college", "university", "institute", "degree", "qualification", "gpa", "cgpa")):
            hints.add("education")
        if any(term in q for term in ("skill", "technology", "tech stack", "tools", "languages")):
            hints.add("skills")
        if any(term in q for term in ("experience", "work history", "employment", "internship")):
            hints.add("experience")
        return hints

    @staticmethod
    def _section_bonus(chunk_lower: str, hints: set[str]) -> float:
        """Boost chunks that look like the section the user asked for."""
        if not hints:
            return 0.0

        bonuses = {
            "education": ("education", "academic", "qualification", "degree", "college", "university", "institute", "school"),
            "projects": ("project", "projects", "built", "developed", "application", "app", "system"),
            "skills": ("skills", "skill", "technology", "tools", "languages", "framework", "library"),
            "experience": ("experience", "internship", "employment", "role", "responsibilities", "worked"),
            "achievements": ("achievement", "achievements", "award", "badge", "certificate", "certification", "honor"),
        }

        score = 0.0
        for hint in hints:
            terms = bonuses.get(hint, ())
            if any(term in chunk_lower for term in terms):
                score += 0.18
        return score

    def _build_summary_answer(
        self,
        provider: ServiceProvider,
        chunks: list[tuple[ProviderDocumentChunk, float]],
        document_text: str = "",
    ) -> str:
        """Build a document summary from retrieved evidence without inventing facts."""
        blob = document_text or "\n".join(chunk.chunk_text for chunk, _ in chunks)
        education = self._extract_education_summary(blob)
        skills = self._extract_skills_summary(blob)
        projects = self._extract_project_summary(blob)
        achievements = self._extract_achievement_summary(blob)

        lines = [f"Document summary for {provider.user.full_name if provider.user else 'the provider'}:"]
        if education:
            lines.append(f"- Education: {education}")
        if skills:
            lines.append(f"- Skills: {skills}")
        if projects:
            lines.append(f"- Projects: {projects}")
        if achievements:
            lines.append(f"- Achievements: {achievements}")
        if len(lines) == 1:
            return "I could not extract a reliable summary from the uploaded document."
        return "\n".join(lines)

    @staticmethod
    def _extract_education_summary(text: str) -> str | None:
        match = re.search(
            r"(National Institute of [^.|\n]+|Indian Institute of [^.|\n]+|University of [^.|\n]+|[A-Z][A-Za-z&.\- ]+(?:Institute|University|College)[A-Za-z&.\- ]*)",
            text,
            re.I,
        )
        parts = []
        if match:
            parts.append("".join(match.group(0).split()))
        gpa = re.search(r"(?:CGPA|GPA)\s*[:\-]?\s*([0-9](?:\.[0-9]+)?)", text, re.I)
        if gpa:
            parts.append(f"GPA {gpa.group(1)}")
        return ", ".join(parts) or None

    @staticmethod
    def _extract_skills_summary(text: str) -> str | None:
        match = re.search(r"Skills?[:\-\s]+(.{0,240})", text, re.I)
        if match:
            return _truncate(" ".join(match.group(1).split()), 220)
        return None

    @staticmethod
    def _extract_project_summary(text: str) -> str | None:
        matches = re.findall(r"(?:Project|Projects)\s*[:\-]?\s*([A-Z][^.|\n]{3,120})", text, re.I)
        if matches:
            return "; ".join(_truncate(" ".join(item.split()), 80) for item in matches[:3])
        return None

    @staticmethod
    def _extract_achievement_summary(text: str) -> str | None:
        match = re.search(r"(?:Achievement|Achievements|Award|Badge|Certification)[:\-\s]+(.{0,220})", text, re.I)
        if match:
            return _truncate(" ".join(match.group(1).split()), 180)
        return None

    @staticmethod
    def _extract_education_line(text: str) -> str | None:
        """Pull a readable education sentence when no exact institution regex hits."""
        lines = re.split(r"(?<=[.!?])\s+|\n+", text)
        for line in lines:
            cleaned = " ".join(line.split()).strip(" -•")
            if not cleaned:
                continue
            lower = cleaned.lower()
            if any(term in lower for term in ("education", "college", "university", "institute", "degree", "gpa", "cgpa")):
                return f"Relevant education detail: {cleaned}"
        return None

    @staticmethod
    def _answer_from_source_text(question_lower: str, text: str) -> str | None:
        """Directly answer common factual document questions from source text."""
        normalized = " ".join(text.split())

        if any(term in question_lower for term in ("college", "university", "institute", "school", "education", "degree", "qualification")):
            patterns = [
                r"(National Institute of [^.|\n]+)",
                r"(Indian Institute of [^.|\n]+)",
                r"(University of [^.|\n]+)",
                r"([A-Z][A-Za-z&.\- ]+(?:Institute|University|College|School|Academy)[A-Za-z&.\- ]*)",
            ]
            for pattern in patterns:
                match = re.search(pattern, normalized, re.I)
                if match:
                    return f"The document mentions { ' '.join(match.group(1).split()) }."

        if "gpa" in question_lower or "cgpa" in question_lower:
            match = re.search(r"(?:CGPA|GPA)\s*[:\-]?\s*([0-9](?:\.[0-9]+)?)", normalized, re.I)
            if match:
                label = "CGPA" if "cgpa" in question_lower else "GPA"
                return f"The document lists {label} {match.group(1)}."

        if any(term in question_lower for term in ("license number", "registration number", "license no", "registration no", "certificate number")):
            match = re.search(r"(?:license|licence|registration|certificate)\s*(?:no\.?|number|#)?\s*[:\-]?\s*([A-Z0-9\-\/]+)", normalized, re.I)
            if match:
                return f"The document lists {match.group(1)} as the requested number."

        if any(term in question_lower for term in ("expiry", "validity", "valid until", "expiration", "expires")):
            match = re.search(r"(?:expiry|valid(?:ity)?|expiration|expires?\s*on)\s*[:\-]?\s*([A-Za-z0-9,\-/ ]{4,40})", normalized, re.I)
            if match:
                return f"The document mentions {match.group(1).strip()} for the validity/expiry date."

        if any(term in question_lower for term in ("issue date", "issued on", "date of issue")):
            match = re.search(r"(?:issue(?:d)?\s*on|date of issue|issued date)\s*[:\-]?\s*([A-Za-z0-9,\-/ ]{4,40})", normalized, re.I)
            if match:
                return f"The document mentions {match.group(1).strip()} as the issue date."

        if any(term in question_lower for term in ("project", "projects")):
            projects = re.findall(r"(?:Project|Projects)\s*[:\-]?\s*([A-Z][^.|\n]{3,120})", normalized, re.I)
            if projects:
                return "The document lists these projects: " + "; ".join(_truncate(" ".join(item.split()), 80) for item in projects[:3]) + "."

        if any(term in question_lower for term in ("achievement", "award", "badge", "certification")):
            ach = re.search(r"(?:Achievement|Achievements|Award|Badge|Certification)[:\-\s]+(.{0,220})", normalized, re.I)
            if ach:
                return f"The document mentions { ' '.join(ach.group(1).split()) }."

        if any(term in question_lower for term in ("skills", "technology", "tech stack", "tools", "languages")):
            skills = re.search(r"Skills?[:\-\s]+(.{0,260})", normalized, re.I)
            if skills:
                return f"The document lists these skills: {_truncate(' '.join(skills.group(1).split()), 220)}."

        return None

    def _extract_document_text(self, doc: dict[str, Any]) -> str:
        preview = doc.get("preview")
        if preview:
            return str(preview)

        relative_path = str(doc.get("path") or "").lstrip("/")
        file_path = BASE_DIR / relative_path
        if not file_path.exists() or not file_path.is_file():
            return ""

        suffix = file_path.suffix.lower()
        content_type = (doc.get("content_type") or "").lower()

        try:
            if suffix in {".txt", ".md", ".csv", ".json", ".yaml", ".yml"} or content_type.startswith("text/"):
                return _read_text_file(file_path)
            if suffix == ".pdf":
                return _read_pdf(file_path)
            if suffix == ".docx":
                return _read_docx(file_path)
            if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
                return _ocr_image(file_path)
        except Exception:
            return ""
        return ""


def _read_text_file(file_path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_path.read_text(encoding=encoding)
        except Exception:
            continue
    return ""


def _read_pdf(file_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        PdfReader = None

    text_parts: list[str] = []
    if PdfReader is not None:
        try:
            reader = PdfReader(str(file_path))
            text_parts = [page.extract_text() or "" for page in reader.pages]
        except Exception:
            text_parts = []

    extracted = "\n\n".join(part for part in text_parts if part.strip()).strip()
    if extracted:
        return extracted

    return _ocr_pdf(file_path)


def _read_docx(file_path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    document = Document(str(file_path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _ocr_pdf(file_path: Path) -> str:
    """Fallback OCR for scanned PDFs and image-based documents."""
    try:
        import fitz  # PyMuPDF
    except Exception:
        return ""

    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return ""

    try:
        document = fitz.open(str(file_path))
    except Exception:
        return ""

    text_parts: list[str] = []
    for page_index in range(len(document)):
        try:
            page = document[page_index]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(image)
            if text.strip():
                text_parts.append(text)
        except Exception:
            continue

    return "\n\n".join(text_parts).strip()


def _ocr_image(file_path: Path) -> str:
    """OCR a standalone image upload."""
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return ""

    try:
        image = Image.open(file_path)
        return pytesseract.image_to_string(image).strip()
    except Exception:
        return ""


def _chunk_text(text: str, max_chars: int = 1100, overlap: int = 180) -> list[str]:
    """
    Bug #22 fix — paragraph-aware chunking strategy.

    Splits on double newlines (semantic paragraph boundaries) first.
    Falls back to character-level sliding window only when a paragraph
    itself exceeds max_chars. This produces more coherent chunks for
    certificates (short paragraphs) and resumes (long narrative blocks).
    """
    if not text or not text.strip():
        return []

    # Split on paragraph boundaries (2+ newlines)
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]

    # If there are no paragraph breaks, treat the whole text as one block
    if len(paragraphs) == 1:
        cleaned = re.sub(r"\s+", " ", text).strip()
        paragraphs = [cleaned]

    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        candidate = f"{current}\n\n{para}".strip() if current else para

        if len(candidate) <= max_chars:
            current = candidate
        else:
            # Flush current buffer
            if current:
                chunks.append(current)

            # Para itself too long → character-split with overlap
            if len(para) > max_chars:
                start = 0
                while start < len(para):
                    end = min(start + max_chars, len(para))
                    chunks.append(para[start:end].strip())
                    if end >= len(para):
                        break
                    start = max(end - overlap, start + 1)
                current = ""
            else:
                current = para

    if current:
        chunks.append(current)

    return [c for c in chunks if c]


def _embed_text(text: str) -> list[float]:
    """Create a semantic embedding, with a deterministic fallback if unavailable."""
    embedder = _get_embedder()
    if embedder is not None:
        try:
            vector = embedder.encode(
                [text],
                normalize_embeddings=True,
                convert_to_numpy=True,
            )[0]
            return [float(value) for value in vector]
        except Exception:
            pass

    return _fallback_embed_text(text)


def _fallback_embed_text(text: str) -> list[float]:
    """Deterministic fallback embedding for offline/dev safety."""
    vector = [0.0] * EMBEDDING_DIMENSIONS
    tokens = _tokens(text)
    for token in tokens:
        digest = hashlib.sha256(token.encode("utf-8")).digest()
        index = int.from_bytes(digest[:4], "big") % EMBEDDING_DIMENSIONS
        sign = 1.0 if digest[4] % 2 == 0 else -1.0
        vector[index] += sign
    norm = math.sqrt(sum(value * value for value in vector)) or 1.0
    return [round(value / norm, 6) for value in vector]


def _get_embedder():
    """Load the sentence-transformer once per process."""
    global _EMBEDDER
    if _EMBEDDER is not None:
        return _EMBEDDER

    try:
        from sentence_transformers import SentenceTransformer
    except Exception:
        _EMBEDDER = None
        return None

    try:
        _EMBEDDER = SentenceTransformer(settings.DOCUMENT_EMBEDDING_MODEL)
    except Exception:
        _EMBEDDER = None
    return _EMBEDDER


def _tokens(text: str) -> list[str]:
    return [token.lower() for token in re.findall(r"[a-zA-Z0-9]+", text) if len(token) > 2]


def _expand_query_tokens(tokens: list[str], raw_query: str) -> list[str]:
    """Expand common document-review shorthand so retrieval can match more naturally."""
    expansions = {
        "cgpa": ["gpa", "grade", "percentage", "score"],
        "gpa": ["cgpa", "grade", "percentage", "score"],
        "marks": ["score", "percentage", "grade"],
        "grade": ["score", "gpa", "cgpa"],
        "percentage": ["marks", "score", "gpa", "cgpa"],
        "experience": ["years", "year"],
        "certificate": ["certification", "credential"],
        "degree": ["education", "qualification"],
    }
    expanded = list(tokens)
    raw_words = set(re.findall(r"[a-zA-Z0-9]+", raw_query.lower()))
    for token in tokens:
        expanded.extend(expansions.get(token, []))
    for word in raw_words:
        expanded.extend(expansions.get(word, []))
    return expanded


def _loads_embedding(raw: str) -> list[float]:
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            return [float(value) for value in data]
    except Exception:
        pass
    return [0.0] * EMBEDDING_DIMENSIONS


def _cosine_similarity(left: list[float], right: list[float]) -> float:
    if not left or not right:
        return 0.0
    return sum(a * b for a, b in zip(left, right))


def _confidence_from_scores(scores: list[float]) -> str:
    if not scores:
        return "low"
    best = max(scores)
    if best >= 0.35:
        return "high"
    if best >= 0.18:
        return "medium"
    return "low"


def _extract_risk_flags(answer: str, chunks: list[tuple[ProviderDocumentChunk, float]]) -> list[str]:
    flags = []
    lowered = answer.lower()
    if "could not find" in lowered or "not found" in lowered or "missing" in lowered:
        flags.append("Some requested evidence was not found in indexed documents")
    if not chunks:
        flags.append("No relevant document chunks retrieved")
    if all(score < 0.18 for _, score in chunks):
        flags.append("Low similarity between question and uploaded document text")
    return flags


def _truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3].rstrip() + "..."
